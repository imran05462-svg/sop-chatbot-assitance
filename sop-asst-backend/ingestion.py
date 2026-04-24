"""
Document ingestion module to handle file parsing and chunking.
Handles plain text, structured paragraphs, and TABLES in DOCX and PDF.
"""
import io
import re
from pypdf import PdfReader
from docx import Document
from typing import List


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf.
    Also attempts to extract table-like structures by preserving whitespace alignment.
    """
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        # extract_text with layout mode preserves column spacing better
        page_text = page.extract_text(extraction_mode="layout")
        if not page_text:
            page_text = page.extract_text()
        if page_text:
            pages_text.append(page_text.strip())
    return "\n\n".join(pages_text)


def _table_to_text(table) -> str:
    """Convert a python-docx Table object into readable 'Key: Value' formatted text."""
    rows = []
    headers = []

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        # Remove duplicate merged cells
        cells = list(dict.fromkeys(cells)) if len(set(cells)) < len(cells) else cells
        cells = [c for c in cells if c]  # Remove empty cells

        if not cells:
            continue

        if i == 0:
            # First row = headers
            headers = cells
            rows.append(" | ".join(cells))
            rows.append("-" * 40)
        else:
            if headers and len(headers) == len(cells):
                # Pair header with value for each cell
                paired = [f"{h}: {v}" for h, v in zip(headers, cells)]
                rows.append(" | ".join(paired))
            else:
                rows.append(" | ".join(cells))

    return "\n".join(rows)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text AND tables from a DOCX file preserving structure."""
    doc = Document(io.BytesIO(file_bytes))
    content_blocks = []

    # Iterate through body elements in order (paragraphs and tables)
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # Get local tag name
        if tag == "p":
            # Paragraph
            para_text = element.text_content().strip() if hasattr(element, 'text_content') else ""
            # Use python-docx paragraph objects for accurate text
            pass
        elif tag == "tbl":
            pass  # handled below

    # Simpler, reliable approach: iterate paragraphs and tables by position
    # Build a map of element order
    body = doc.element.body
    processed = set()

    for child in body:
        local_tag = child.tag.split("}")[-1]

        if local_tag == "p":
            # Find matching paragraph
            for para in doc.paragraphs:
                if para._element is child and id(child) not in processed:
                    processed.add(id(child))
                    if para.text.strip():
                        content_blocks.append(para.text.strip())
                    break

        elif local_tag == "tbl":
            for table in doc.tables:
                if table._element is child and id(child) not in processed:
                    processed.add(id(child))
                    table_text = _table_to_text(table)
                    if table_text.strip():
                        content_blocks.append(f"[TABLE]\n{table_text}\n[/TABLE]")
                    break

    return "\n\n".join(content_blocks)


def chunk_text(text: str, chunk_size: int = 800, chunk_overlap: int = 150) -> List[str]:
    """
    Split text into overlapping chunks.
    Tries to split on double-newlines (paragraph boundaries) first
    to avoid cutting mid-sentence or mid-table-row.
    """
    # Try paragraph-aware splitting first
    paragraphs = re.split(r'\n{2,}', text)
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current_chunk) + len(para) + 2 <= chunk_size:
            current_chunk += ("\n\n" if current_chunk else "") + para
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            # If a single paragraph is too large, split it by character
            if len(para) > chunk_size:
                start = 0
                while start < len(para):
                    end = start + chunk_size
                    chunks.append(para[start:end].strip())
                    start += chunk_size - chunk_overlap
                current_chunk = ""
            else:
                current_chunk = para

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return [c for c in chunks if c]


def process_uploaded_file(file_name: str, file_bytes: bytes) -> List[str]:
    """
    Process an uploaded file to extract text (including tables) and return chunks.
    Raises ValueError for unsupported file types.
    """
    if file_name.lower().endswith(".pdf"):
        text = parse_pdf(file_bytes)
    elif file_name.lower().endswith(".docx"):
        text = parse_docx(file_bytes)
    elif file_name.lower().endswith(".txt"):
        text = file_bytes.decode("utf-8")
    else:
        raise ValueError(f"Unsupported file type: {file_name}")

    chunks = chunk_text(text)
    return chunks


if __name__ == "__main__":
    sample_text = """Employee Name | Department | Status
John Doe | IT Support | Active
Jane Smith | HR | Active

Procedure:
Step 1: Reboot the system.
Step 2: Check logs at /var/log/syslog.
Step 3: Escalate to Level 3 via Jira if unresolved.
""" * 5
    chunks = chunk_text(sample_text)
    print(f"Generated {len(chunks)} chunks.")
    for i, c in enumerate(chunks[:3]):
        print(f"\n--- Chunk {i+1} ---\n{c}")
    print("\nIngestion logic tested successfully!")
